"""Build the visually verified QLoRA check/evaluation guide."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "FINETUNED_MODEL_CHECK_AND_EVALUATION.md"
OUTPUT = HERE / "FINETUNED_MODEL_CHECK_AND_EVALUATION.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.20

    header = section.header.paragraphs[0]
    header.text = "THG Lab | 1M-row QLoRA evaluation guide"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    add_field(footer, "PAGE")


def inline(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_BLUE
        else:
            paragraph.add_run(part)


def title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(24)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("LAWRENCIUM MODEL HANDOFF")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run("Checking and Evaluating the 1M-row QLoRA Model")
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(20, 42, 68)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        "Verified completion evidence, exact loading steps, a reproducible "
        "base-model comparison, and publication-safe next steps"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    for label, value in (
        ("Training result", "COMPLETED - 12,090/12,090 steps - exit 0"),
        ("Hardware", "Four NVIDIA A40 GPUs"),
        ("Dataset", "1,000,000 rows - one epoch"),
        ("Updated", "August 31, 2026"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)


def build() -> None:
    doc = Document()
    configure(doc)
    title_block(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("# ") or line.startswith("**Status:") or line.startswith("**SLURM result:") or line.startswith("**Adapter:"):
            index += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.08)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_together = True
                run = p.add_run("\n".join(code))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(38, 50, 56)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                p._p.get_or_add_pPr().append(shading)
                in_code = False
            index += 1
            continue
        if in_code:
            code.append(line)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            inline(p, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline(p, line[2:])
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not re.match(
            r"^(#|```|- |\d+\. )", lines[index]
        ):
            paragraph_lines.append(lines[index])
            index += 1
        p = doc.add_paragraph()
        inline(p, " ".join(paragraph_lines))

    doc.core_properties.title = "Checking and Evaluating the 1M-row QLoRA Model"
    doc.core_properties.subject = "Lawrencium QLoRA completion and evaluation runbook"
    doc.core_properties.author = "THG Lab project team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
