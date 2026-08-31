"""Build the PI-facing SynAgent/SynLlama/fine-tuning report and progress plot."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PI_SYNAGENT_SYNLLAMA_AND_FINETUNING_REPORT.md"
OUTPUT = ROOT / "docs" / "PI_SYNAGENT_SYNLLAMA_AND_FINETUNING_REPORT.docx"
PROGRESS = ROOT / "docs" / "chembl-benchmark" / "figures" / "fig5_finetuning_progress.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)
LIGHT = "F2F4F7"


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(total))
    width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    with Image.open(path) as image:
        width_px, height_px = image.size
    max_w, max_h = 6.15, 7.0
    ratio = width_px / height_px
    width = min(max_w, max_h * ratio)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def make_progress_figure() -> None:
    import matplotlib.pyplot as plt

    steps = list(range(500, 12001, 500))
    losses = [
        0.10089174, 0.08375327, 0.07617992, 0.07127548, 0.06849218,
        0.06579643, 0.06446034, 0.06168590, 0.06022490, 0.05911352,
        0.05816337, 0.05668017, 0.05565141, 0.05455855, 0.05353912,
        0.05278006, 0.05188552, 0.05125288, 0.05064403, 0.05009909,
        0.04969271, 0.04946361, 0.04941533, 0.04937038,
    ]
    plt.style.use("seaborn-v0_8-whitegrid")
    fig, ax = plt.subplots(figsize=(8.6, 4.6))
    ax.plot(steps, losses, color="#2E74B5", linewidth=2.6, marker="o", markersize=4)
    ax.fill_between(steps, losses, min(losses) - 0.003, color="#2E74B5", alpha=0.10)
    ax.set_title("1M-row QLoRA held-out loss through step 12,000", loc="left", weight="bold")
    ax.set_xlabel("Optimizer step")
    ax.set_ylabel("Evaluation loss")
    ax.annotate("0.1009", (steps[0], losses[0]), xytext=(8, 4), textcoords="offset points")
    ax.annotate("0.0494", (steps[-1], losses[-1]), xytext=(-36, 8), textcoords="offset points")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(PROGRESS, dpi=180, bbox_inches="tight")
    plt.close(fig)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "THG Lab | SynAgent research status"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    add_field(footer, "PAGE")


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("SYNAGENT RESEARCH STATUS")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("SynLlama comparison, agent correction evidence, and 1M-row QLoRA training")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    for label, value in (
        ("Prepared for", "THG Lab project discussion"),
        ("Updated", "August 31, 2026"),
        ("Branch", "lukasmki/SynAgent: synagent-full-pipeline"),
        ("Training", "Lawrencium job 25306635 - completed 12,090/12,090, exit 0"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
    rule = doc.add_paragraph()
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


FIGURES = {
    "3. Complete 10,000-path comparison": (
        ROOT / "docs/chembl-benchmark/comparison-2026-08-27/synagent-vs-synllama-summary.png",
        "Figure 1. Full 10,000-path validator comparison and paired n=50 correction pilot.",
    ),
    "4. Genuine agent-correction pilot": (
        ROOT / "docs/chembl-benchmark/comparison-2026-08-27/three-synagent-wins.png",
        "Figure 2. Three representative fail-to-pass corrections from the genuine agent batch.",
    ),
    "7. Real 1M-row QLoRA run": (
        PROGRESS,
        "Figure 3. Held-out evaluation loss decreased monotonically from step 500 through step 12,000.",
    ),
}


def build() -> None:
    if "--no-chart" not in sys.argv:
        make_progress_figure()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        line = lines[idx]
        if line.startswith("# "):
            idx += 1
            continue
        if line.startswith("**Prepared for:") or line.startswith("**Updated:") or line.startswith("**Repository scope:"):
            idx += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                in_code = False
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if line.startswith("## "):
            heading = line[3:]
            doc.add_heading(heading, level=1)
            if heading in FIGURES:
                add_figure(doc, *FIGURES[heading])
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            idx += 1
            continue
        if line.startswith("| ") and idx + 1 < len(lines) and lines[idx + 1].startswith("|---"):
            rows: list[list[str]] = []
            header = [c.strip() for c in line.strip("|").split("|")]
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                rows.append([c.strip() for c in lines[idx].strip("|").split("|")])
                idx += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for col, value in enumerate(header):
                table.rows[0].cells[col].text = value
                shade(table.rows[0].cells[col], LIGHT)
                for run in table.rows[0].cells[col].paragraphs[0].runs:
                    run.bold = True
            for values in rows:
                cells = table.add_row().cells
                for col, value in enumerate(values):
                    cells[col].text = re.sub(r"\*\*", "", value)
                    cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            widths = [round(9360 / len(header))] * len(header)
            widths[-1] += 9360 - sum(widths)
            set_table_geometry(table, widths)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if re.match(r"^[1-9]\. ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[1-9]\. ", "", line))
            idx += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            idx += 1
            continue
        if line.startswith("> "):
            quote = []
            while idx < len(lines) and (lines[idx].startswith("> ") or not lines[idx].strip()):
                if lines[idx].startswith("> "):
                    quote.append(lines[idx][2:])
                idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            run = p.add_run(" ".join(quote))
            run.italic = True
            run.font.color.rgb = DARK_BLUE
            continue
        if not line.strip():
            idx += 1
            continue
        paragraph_lines = [line]
        idx += 1
        while idx < len(lines) and lines[idx].strip() and not re.match(r"^(#|\||- |[1-9]\. |>|```)", lines[idx]):
            paragraph_lines.append(lines[idx])
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))

    doc.core_properties.title = "SynAgent Research Status"
    doc.core_properties.subject = "SynLlama comparison and Lawrencium fine-tuning"
    doc.core_properties.author = "THG Lab project team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
